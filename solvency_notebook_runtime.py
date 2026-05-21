"""Runtime backend for the Solvency II Streamlit app.

This file is intentionally static Python code. It replaces the old notebook
loader so the public app does not rely on dynamic notebook execution or process-wide directory changes.
"""
# ---- Notebook cell 6 ----
from __future__ import annotations

import hashlib
import json
import os
import re
import unicodedata
from datetime import datetime, timezone
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, Optional

import pandas as pd
from bs4 import BeautifulSoup
from rank_bm25 import BM25Okapi

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import chromadb
    from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
except Exception as e:
    chromadb = None
    SentenceTransformerEmbeddingFunction = None
    print(f"Chroma embedding support unavailable: {type(e).__name__}: {e}")

try:
    from sentence_transformers import CrossEncoder
except Exception as e:
    CrossEncoder = None
    print(f"CrossEncoder unavailable: {type(e).__name__}: {e}")


# ---- Notebook cell 7 ----
# Configuration

PROJECT_DIR = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
DOCS_DIR = PROJECT_DIR / "Directive"
INDEX_DIR = PROJECT_DIR / "rag3_index"
CHUNKS_PATH = INDEX_DIR / "chunks.jsonl"
MANIFEST_PATH = INDEX_DIR / "manifest.json"
AUDIT_LOG_PATH = PROJECT_DIR / "audit_log.jsonl"

# Les modèles multilingues sont plus sûrs pour les contenus réglementaires en français que les configurations par défaut limitées à l’anglais.
EMBEDDING_MODEL = "intfloat/multilingual-e5-large"
RERANKER_MODEL = "cross-encoder/mmarco-mMiniLMv2-L12-H384-v1"

TOP_K_BM25 = 30
TOP_K_VECTOR = 30
TOP_K_FINAL = 4
RRF_K = 30
BM25_RRF_WEIGHT = 0.60
VECTOR_RRF_WEIGHT = 0.40
MIN_EXTRACTED_TEXT_CHARS = 50
MIN_RETRIEVED_CHARS = 200
DOCUMENT_TYPE_BOOSTS = {
    "Directive Solvabilité II": 1.20,
    "Delegated Regulation": 1.12,
    "EIOPA material": 1.05,
    "Superviseur local": 1.08,
    "Matériel ORSA/ERSA": 1.05,
    "Matériel SFCR": 1.03,
    "Matériel RSR": 1.03,
    "Inconnu": 1.00,
}

DOCS_DIR.mkdir(exist_ok=True)
INDEX_DIR.mkdir(exist_ok=True)

print(f"Corpus Solvabilité II embarqué: {DOCS_DIR.resolve()}")


# ---- Notebook cell 9 ----
# Structure les documents et les morceaux de texte pour que chaque information utilisée par le RAG
# garde son origine : fichier source, page, type de document, version, juridiction, etc.
# Cela permet de retrouver exactement d'où vient une réponse, de citer les sources,
# d'éviter les doublons et de rendre le système vérifiable.
# Définit les objets utilisés pour stocker un document complet et ses morceaux de texte.
# Chaque document ou chunk garde son contenu, sa source et ses métadonnées
# afin de savoir précisément d'où vient l'information utilisée par le RAG.
# La fonction stable_id crée un identifiant stable pour retrouver facilement
# le même document ou le même chunk à chaque exécution.
@dataclass
class LoadedDocument:
    text: str
    source_path: str
    source_name: str
    page: Optional[int] = None
    document_type: Optional[str] = None
    jurisdiction: Optional[str] = None
    version: Optional[str] = None
    temporal_note: Optional[str] = None


@dataclass
class Chunk:
    chunk_id: str
    text: str
    source_path: str
    source_name: str
    page: Optional[int]
    section: Optional[str]
    document_type: Optional[str]
    jurisdiction: Optional[str]
    version: Optional[str]
    temporal_note: Optional[str] = None


def stable_id(*parts: str) -> str:
    raw = "||".join(str(part) for part in parts)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:24]


# ---- Notebook cell 11 ----
# Analyse le nom d'un fichier pour déduire automatiquement ses métadonnées principales :
# type de document, juridiction, version/date et note de contexte.
# Ces informations permettent au RAG de mieux classer les sources, de distinguer les textes
# européens, locaux ou EIOPA, et de garder une trace du contexte réglementaire applicable.

def infer_metadata(path: Path) -> dict:
    name = path.stem.lower()

    if "eiopa" in name:
        document_type = "EIOPA material"
    elif "delegated" in name or "2015-35" in name:
        document_type = "Delegated Regulation"
    elif (
        "directive" in name
        or "2009-138" in name
        or "32009l0138" in name
        or "32025l0002" in name
        or "202500002" in name
        or "celex" in name
    ):
        document_type = "Directive Solvabilité II"
    elif "orsa" in name or "ersa" in name:
        document_type = "Matériel ORSA/ERSA"
    elif "sfcr" in name or "rapport sur la solvabilite" in name:
        document_type = "Matériel SFCR"
    elif "rsr" in name:
        document_type = "Matériel RSR"
    elif "bnb" in name or "nbb" in name:
        document_type = "Superviseur local"
    else:
        document_type = "Inconnu"

    jurisdiction = "EU" if (
        "celex" in name
        or "32009l0138" in name
        or "32025l0002" in name
        or "202500002" in name
    ) else None
    if "bnb" in name or "nbb" in name:
        jurisdiction = "BE"
    for candidate in ["EU", "BE", "FR", "LU", "NL", "DE", "UK"]:
        if re.search(rf"(^|[_\- ]){candidate.lower()}($|[_\- ])", name):
            jurisdiction = candidate
            break

    version_match = re.search(r"(20\d{2}[-_ ]?\d{2}[-_ ]?\d{2}|v\d+(?:\.\d+)?)", name)
    version = version_match.group(1) if version_match else None

    if document_type == "Directive Solvabilité II":
        if "32025l0002" in name or "202500002" in name:
            version = "Directive (UE) 2025/2"
            temporal_note = (
                "Directive (UE) 2025/2 du 27/11/2024, publiée au JOUE le 08/01/2025; "
                "transposition au plus tard le 29/01/2027 et application à partir du 30/01/2027."
            )
        else:
            version = version or "2009/138/CE"
            temporal_note = (
                "Directive 2009/138/CE Solvabilité II; cadre applicable depuis le 01/01/2016, "
                "tel que modifié notamment par Omnibus II et sous réserve des dispositions transitoires."
            )
    elif document_type == "Delegated Regulation":
        version = version or "2015/35"
        temporal_note = (
            "Règlement délégué (UE) 2015/35 complétant Solvabilité II; "
            "à lire avec la Directive 2009/138/CE et ses modifications ultérieures."
        )
    elif document_type == "Superviseur local":
        version = version or "NBB_2016_31"
        temporal_note = (
            "Source de superviseur local belge; vérifier la date/version de la circulaire "
            "et son éventuelle consolidation avant usage prudentiel."
        )
    elif document_type == "EIOPA material":
        temporal_note = (
            "Orientation ou Q&A EIOPA; clarification de niveau 3. Vérifier la date de publication "
            "et les éventuelles mises à jour sur la page source."
        )
    else:
        temporal_note = None

    return {
        "document_type": document_type,
        "jurisdiction": jurisdiction,
        "version": version,
        "temporal_note": temporal_note,
    }


# ---- Notebook cell 13 ----
# Charge les documents réglementaires depuis différents formats (PDF, DOCX, TXT, MD, HTML)
# et les transforme en objets LoadedDocument enrichis avec leurs métadonnées.
# Pour les PDF, le code extrait le texte page par page avec PyMuPDF, puis utilise l'OCR
# lorsque le texte extrait est presque vide, ce qui permet de traiter aussi certains PDF scannés.
# Les fichiers Word, texte, Markdown et HTML sont lus avec le chargeur adapté,
# puis tous les documents trouvés dans le dossier sont parcourus et chargés automatiquement.
def _pymupdf_ocr_text(page) -> str:
    """OCR a page with the conda Tesseract install when available."""
    if fitz is None:
        return ""

    env_bin = Path(os.environ.get("CONDA_PREFIX", Path(os.sys.executable).resolve().parent.parent)) / "bin"
    tessdata = env_bin.parent / "share" / "tessdata"
    os.environ["PATH"] = str(env_bin) + os.pathsep + os.environ.get("PATH", "")

    try:
        kwargs = {"language": "eng", "dpi": 200, "full": True}
        if tessdata.exists():
            kwargs["tessdata"] = str(tessdata)
        textpage = page.get_textpage_ocr(**kwargs)
        return page.get_text("text", textpage=textpage) or ""
    except Exception:
        return ""


def load_pdf(path: Path) -> list[LoadedDocument]:
    """Load PDF pages with PyMuPDF first, with OCR fallback for near-empty pages."""
    metadata = infer_metadata(path)
    docs = []

    if fitz is not None:
        pdf = fitz.open(str(path))
        try:
            for page_index, page in enumerate(pdf, start=1):
                text = page.get_text("text") or ""
                if len(text.strip()) < MIN_EXTRACTED_TEXT_CHARS:
                    ocr_text = _pymupdf_ocr_text(page)
                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                if text.strip():
                    docs.append(LoadedDocument(
                        text=text,
                        source_path=str(path),
                        source_name=path.name,
                        page=page_index,
                        **metadata,
                    ))
        finally:
            pdf.close()
        return docs

    if PdfReader is None:
        raise ImportError("Install pymupdf to load PDFs robustly: pip install pymupdf")

    reader = PdfReader(str(path))
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(LoadedDocument(
                text=text,
                source_path=str(path),
                source_name=path.name,
                page=page_index,
                **metadata,
            ))
    return docs


def load_docx(path: Path) -> list[LoadedDocument]:
    if docx is None:
        raise ImportError("Install python-docx to load DOCX files: pip install python-docx")

    metadata = infer_metadata(path)
    document = docx.Document(str(path))
    text = "\n\n".join(p.text for p in document.paragraphs if p.text.strip())
    return [LoadedDocument(text=text, source_path=str(path), source_name=path.name, **metadata)]


def load_text_like(path: Path) -> list[LoadedDocument]:
    metadata = infer_metadata(path)
    raw = path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix.lower() in {".html", ".htm"}:
        soup = BeautifulSoup(raw, "html.parser")
        raw = soup.get_text("\n", strip=True)
    return [LoadedDocument(text=raw, source_path=str(path), source_name=path.name, **metadata)]


def load_documents(folder: Path) -> list[LoadedDocument]:
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_text_like,
        ".md": load_text_like,
        ".html": load_text_like,
        ".htm": load_text_like,
    }

    documents = []
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in loaders:
            loaded = loaders[path.suffix.lower()](path)
            documents.extend(loaded)
            print(f"Loaded {len(loaded):>3} item(s): {path.name}")

    return documents


# ---- Notebook cell 15 ----
SECTION_RE = re.compile(
    r"""(?imx)
    ^
    (
        article\s+\d+[a-z]*(?:\s+bis|\s+ter)?
        | considérant\s+\(?\d+\)?
        | orientation\s+\d+
        | ligne\s+directrice\s+\d+
        | guideline\s+\d+
        | chapitre\s+[ivxlcdm\d]+
        | chapter\s+[ivxlcdm\d]+
        | titre\s+[ivxlcdm\d]+
        | title\s+[ivxlcdm\d]+
        | annexe\s+[ivxlcdm\d]+
        | annex\s+[ivxlcdm\d]+
        | section\s+\d+
        | \d+\.\s+[A-ZÉÈÀÂÊÎÔÛÇ][^\n]{3,140}
    )
    """,
)

ARTICLE_SECTION_RE = re.compile(r"(?i)^article\s+\d+[a-z]*(?:\s+bis|\s+ter)?$")
ARTICLE_MAX_CHARS = 7000
DEFAULT_CHUNK_MAX_CHARS = 2400
CHUNK_OVERLAP_CHARS = 350
CHUNKING_VERSION = "article-aware-v2"


def normalize_whitespace(text: str) -> str:
    text = re.sub(r"\u00a0", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_by_sections_with_offsets(text: str) -> list[tuple[Optional[str], str, int]]:
    text = normalize_whitespace(text)
    matches = list(SECTION_RE.finditer(text))

    # In legal texts, numbered paragraphs such as "1." and "2." belong to the
    # surrounding article; they should not split Article 77 across pages.
    if any(ARTICLE_SECTION_RE.match(re.sub(r"\s+", " ", match.group(1).strip())) for match in matches):
        matches = [
            match for match in matches
            if not re.match(r"^\d+\.", re.sub(r"\s+", " ", match.group(1).strip()))
        ]

    if not matches:
        return [(None, text, 0)] if text else []

    sections = []
    if matches[0].start() > 0:
        prefix = text[:matches[0].start()].strip()
        if prefix:
            sections.append((None, prefix, 0))

    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section_name = re.sub(r"\s+", " ", match.group(1).strip())
        section_text = text[start:end].strip()
        if section_text:
            sections.append((section_name, section_text, start))

    return sections


def split_by_sections(text: str) -> list[tuple[Optional[str], str]]:
    return [(section, section_text) for section, section_text, _ in split_by_sections_with_offsets(text)]


NUMBERED_PARAGRAPH_RE = re.compile(
    r"(?m)^(?:\d+\.|\(\d+\)|[a-z]\))\s+"
)


def split_numbered_paragraphs(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []

    matches = list(NUMBERED_PARAGRAPH_RE.finditer(text))
    if not matches:
        return [text]

    blocks = []
    if matches[0].start() > 0:
        prefix = text[:matches[0].start()].strip()
        if prefix:
            blocks.append(prefix)

    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        block = text[start:end].strip()
        if block:
            blocks.append(block)

    return blocks


def window_text(text: str, max_chars: int = DEFAULT_CHUNK_MAX_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    text = text.strip()
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    atomic_blocks = split_numbered_paragraphs(text)
    if len(atomic_blocks) > 1:
        windows = []
        current = ""

        for block in atomic_blocks:
            if len(block) > max_chars:
                if current:
                    windows.append(current.strip())
                    current = ""
                windows.extend(window_text(block, max_chars=max_chars, overlap=overlap))
                continue

            candidate = f"{current}\n\n{block}".strip() if current else block
            if len(candidate) <= max_chars:
                current = candidate
            else:
                if current:
                    windows.append(current.strip())
                current = block

        if current:
            windows.append(current.strip())
        return [w for w in windows if w]

    windows = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        cut = text[start:end]

        if end < len(text):
            paragraph_cut = cut.rfind("\n\n")
            sentence_cut = max(cut.rfind(". "), cut.rfind("; "), cut.rfind(": "))
            boundary = paragraph_cut if paragraph_cut > max_chars * 0.55 else sentence_cut
            if boundary > max_chars * 0.55:
                cut = cut[: boundary + 1]
                end = start + boundary + 1

        windows.append(cut.strip())
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)

    return [w for w in windows if w]


def _group_documents_by_source(documents: list[LoadedDocument]) -> list[dict]:
    grouped = []
    by_source = {}
    for doc in documents:
        group = by_source.get(doc.source_path)
        if group is None:
            group = {
                "source_path": doc.source_path,
                "source_name": doc.source_name,
                "document_type": doc.document_type,
                "jurisdiction": doc.jurisdiction,
                "version": doc.version,
                "temporal_note": doc.temporal_note,
                "docs": [],
            }
            by_source[doc.source_path] = group
            grouped.append(group)
        group["docs"].append(doc)

    for group in grouped:
        group["docs"].sort(key=lambda doc: (doc.page is None, doc.page or 0))
    return grouped


def _combine_group_pages(group: dict) -> tuple[str, list[tuple[int, Optional[int]]]]:
    combined_parts = []
    page_offsets = []
    cursor = 0
    for doc in group["docs"]:
        text = normalize_whitespace(doc.text)
        if not text:
            continue
        if combined_parts:
            separator = "\n\n"
            combined_parts.append(separator)
            cursor += len(separator)
        page_offsets.append((cursor, doc.page))
        combined_parts.append(text)
        cursor += len(text)
    return "".join(combined_parts), page_offsets


def _page_for_offset(page_offsets: list[tuple[int, Optional[int]]], offset: int) -> Optional[int]:
    current_page = None
    for page_start, page in page_offsets:
        if page_start <= offset:
            current_page = page
        else:
            break
    return current_page


def _max_chars_for_section(section: Optional[str], section_text: str) -> int:
    if section and ARTICLE_SECTION_RE.match(section.strip()):
        return max(ARTICLE_MAX_CHARS, min(len(section_text), ARTICLE_MAX_CHARS))
    return DEFAULT_CHUNK_MAX_CHARS


def chunk_documents(documents: list[LoadedDocument]) -> list[Chunk]:
    chunks = []
    for group in _group_documents_by_source(documents):
        combined_text, page_offsets = _combine_group_pages(group)
        if not combined_text.strip():
            continue

        for section_index, (section, section_text, section_offset) in enumerate(split_by_sections_with_offsets(combined_text), start=1):
            max_chars = _max_chars_for_section(section, section_text)
            parts = window_text(section_text, max_chars=max_chars, overlap=CHUNK_OVERLAP_CHARS)
            for part_index, part in enumerate(parts, start=1):
                relative_offset = section_text.find(part[: min(len(part), 80)])
                if relative_offset < 0:
                    relative_offset = 0
                absolute_offset = section_offset + relative_offset
                page = _page_for_offset(page_offsets, absolute_offset)
                chunk_id = stable_id(
                    CHUNKING_VERSION,
                    group["source_path"],
                    str(page),
                    str(section_index),
                    str(section),
                    str(part_index),
                    part,
                )
                chunks.append(Chunk(
                    chunk_id=chunk_id,
                    text=part,
                    source_path=group["source_path"],
                    source_name=group["source_name"],
                    page=page,
                    section=section,
                    document_type=group["document_type"],
                    jurisdiction=group["jurisdiction"],
                    version=group["version"],
                    temporal_note=group["temporal_note"],
                ))
    return chunks


# ---- Notebook cell 17 ----
# Gère la création et la mise à jour de l’index documentaire du RAG.
# Le code sauvegarde les chunks dans un fichier auditable, vérifie si les sources ont changé,
# recharge les chunks existants si tout est déjà à jour, ou reconstruit les chunks si nécessaire.
# Il détecte aussi les pages presque vides pour signaler un possible besoin d’OCR.
# Ensuite, il construit l’index de recherche : un stockage JSONL pour l’audit,
# un index BM25 pour la recherche par mots-clés, et si les bibliothèques sont disponibles,
# un index vectoriel Chroma pour la recherche sémantique.

def save_chunks(chunks: list[Chunk], path: Path) -> None:
    path.parent.mkdir(exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(json.dumps(asdict(chunk), ensure_ascii=False) + "\n")


def load_chunks(path: Path) -> list[Chunk]:
    with path.open("r", encoding="utf-8") as f:
        return [Chunk(**json.loads(line)) for line in f if line.strip()]


def source_manifest(folder: Path) -> dict:
    files = []
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}:
            stat = path.stat()
            files.append({
                "path": str(path),
                "size": stat.st_size,
                "mtime_ns": stat.st_mtime_ns,
            })
    return {
        "embedding_model": EMBEDDING_MODEL,
        "chunking_version": CHUNKING_VERSION,
        "article_max_chars": ARTICLE_MAX_CHARS,
        "default_chunk_max_chars": DEFAULT_CHUNK_MAX_CHARS,
        "files": files,
    }


def chunks_are_current() -> bool:
    if not CHUNKS_PATH.exists() or not MANIFEST_PATH.exists():
        return False
    try:
        return json.loads(MANIFEST_PATH.read_text(encoding="utf-8")) == source_manifest(DOCS_DIR)
    except Exception:
        return False


def vector_index_is_ready(expected_count: Optional[int] = None) -> bool:
    if chromadb is None or SentenceTransformerEmbeddingFunction is None:
        return True

    chroma_dir = INDEX_DIR / "chroma"
    if not chroma_dir.exists():
        return False

    try:
        client = chromadb.PersistentClient(path=str(chroma_dir))
        collection = client.get_collection("solvency_ii_chunks")
        count = collection.count()
        return count > 0 and (expected_count is None or count == expected_count)
    except Exception:
        return False


def index_is_current() -> bool:
    if not chunks_are_current():
        return False
    try:
        chunk_count = sum(1 for line in CHUNKS_PATH.read_text(encoding="utf-8").splitlines() if line.strip())
        return vector_index_is_ready(expected_count=chunk_count)
    except Exception:
        return False


def build_chunks(force: bool = False) -> list[Chunk]:
    if not force and chunks_are_current():
        print("Chunks déjà à jour.")
        return load_chunks(CHUNKS_PATH)

    documents = load_documents(DOCS_DIR)
    if not documents:
        raise ValueError(f"No supported documents found in {DOCS_DIR.resolve()}")

    empty_pages = [doc for doc in documents if len(doc.text.strip()) < MIN_EXTRACTED_TEXT_CHARS]
    if empty_pages:
        print(
            f"Attention: {len(empty_pages)} page(s)/document(s) ignorées car presque vides. "
            "OCR possiblement nécessaire."
        )
        scan_report = [
            {
                "source_name": doc.source_name,
                "source_path": doc.source_path,
                "page": doc.page,
                "extracted_chars": len(doc.text.strip()),
            }
            for doc in empty_pages
        ]
        (INDEX_DIR / "empty_extraction_report.json").write_text(
            json.dumps(scan_report, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    documents = [doc for doc in documents if len(doc.text.strip()) >= MIN_EXTRACTED_TEXT_CHARS]
    if not documents:
        raise ValueError("All documents/pages look empty after extraction. OCR is required before indexing.")

    chunks = chunk_documents(documents)
    if not chunks:
        raise ValueError("No chunks created. Check PDF extraction/OCR.")

    save_chunks(chunks, CHUNKS_PATH)
    print(f"Saved {len(chunks)} chunks to {CHUNKS_PATH}")

    MANIFEST_PATH.write_text(json.dumps(source_manifest(DOCS_DIR), ensure_ascii=False, indent=2), encoding="utf-8")
    return chunks


def build_indexes(force: bool = False) -> list[Chunk]:
    if not force and index_is_current():
        print("Index déjà à jour.")
        return load_chunks(CHUNKS_PATH)

    chunks = build_chunks(force=force)

    if chromadb is None or SentenceTransformerEmbeddingFunction is None:
        print("Index BM25 prêt. Index vectoriel ignoré: chromadb/sentence-transformers indisponible.")
        return chunks

    try:
        embedding_function = SentenceTransformerEmbeddingFunction(model_name=EMBEDDING_MODEL)
        client = chromadb.PersistentClient(path=str(INDEX_DIR / "chroma"))

        try:
            client.delete_collection("solvency_ii_chunks")
        except Exception:
            pass

        collection = client.get_or_create_collection(
            name="solvency_ii_chunks",
            embedding_function=embedding_function,
            metadata={"hnsw:space": "cosine"},
        )

        batch_size = 128
        for start in range(0, len(chunks), batch_size):
            batch = chunks[start:start + batch_size]
            collection.add(
                ids=[chunk.chunk_id for chunk in batch],
                documents=[chunk.text for chunk in batch],
                metadatas=[{
                    "source_name": chunk.source_name,
                    "source_path": chunk.source_path,
                    "page": chunk.page or "",
                    "section": chunk.section or "",
                    "document_type": chunk.document_type or "",
                    "jurisdiction": chunk.jurisdiction or "",
                    "version": chunk.version or "",
                } for chunk in batch],
            )
            print(f"Indexed {min(start + batch_size, len(chunks))}/{len(chunks)} chunks")
    except Exception as e:
        print(f"Index BM25 prêt. Index vectoriel ignoré: {type(e).__name__}: {e}")
        return chunks

    print("Vector index ready")
    return chunks


def ensure_chunks() -> list[Chunk]:
    return build_chunks(force=False)


def ensure_indexes() -> list[Chunk]:
    return build_indexes(force=False)


# Run this before asking questions.
chunks = ensure_chunks()


# ---- Notebook cell 19 ----
FRENCH_STOPWORDS = {
    "le", "la", "les", "un", "une", "des", "de", "du", "d", "l", "et", "ou", "à", "a",
    "au", "aux", "en", "dans", "sur", "pour", "par", "que", "qui", "quoi", "dont",
    "est", "sont", "avec", "ce", "cet", "cette", "ces", "selon", "directive"
}

QUERY_SYNONYMS = {
    "revoyure": [
        "Directive (UE) 2025/2",
        "directive 2025/2",
        "32025L0002",
        "ELI: http://data.europa.eu/eli/dir/2025/2/oj",
        "30 janvier 2027",
        "29 janvier 2027",
        "proportionnalité",
        "entreprises de petite taille et non complexes",
        "marge de risque",
        "volatility adjustment",
        "ajustement de volatilité",
        "mesures de garantie de long terme",
        "risques en matière de durabilité",
        "risque de liquidité",
        "surveillance macroprudentielle",
        "qualité du contrôle",
    ],
    "review": [
        "Directive (EU) 2025/2",
        "Solvency II review",
        "Directive 2025/2",
        "30 January 2027",
        "proportionality",
        "small and non-complex undertakings",
        "risk margin",
        "volatility adjustment",
        "sustainability risks",
        "liquidity risk",
        "macro-prudential supervision",
    ],
    "best estimate": ["meilleure estimation"],
    "be": ["best estimate", "meilleure estimation"],
    "scr": ["capital de solvabilité requis"],
    "mcr": ["minimum de capital requis"],
    "gouvernance": [
        "système de gouvernance",
        "article 40",
        "article 41",
        "article 42",
        "article 43",
        "article 44",
        "article 45",
        "article 46",
        "article 47",
        "article 48",
        "article 49",
        "chapitre iv",
        "exigences qualitatives",
    ],
    "orsa": [
        "article 45",
        "évaluation interne des risques",
        "évaluation interne des risques et de la solvabilité",
        "ersa",
        "own risk solvency assessment",
        "own risk and solvency assessment",
    ],
    "sfcr": ["rapport sur la solvabilité et la situation financière"],
}

SOLVENCY_II_REVIEW_QUERY_TERMS = {
    "revoyure",
    "review",
    "2025/2",
    "2025-2",
    "32025l0002",
    "202500002",
    "réexamen",
    "reexamen",
}

SOLVENCY_II_REVIEW_SOURCE_TERMS = {
    "32025l0002",
    "202500002",
    "oj_l_202500002",
    "directive (ue) 2025/2",
    "directive (eu) 2025/2",
    "eli/dir/2025/2",
}

SOLVENCY_SCOPE_TERMS = {
    "solvabilite ii",
    "solvency ii",
    "s2",
    "assurance",
    "assureur",
    "assureurs",
    "reassurance",
    "prudentiel",
    "prudentielle",
    "actuariel",
    "actuaire",
    "provisions techniques",
    "provision technique",
    "best estimate",
    "meilleure estimation",
    "risk margin",
    "marge de risque",
    "fonds propres",
    "fonds propres eligible",
    "fonds propres eligibles",
    "own funds",
    "basic own funds",
    "ancillary own funds",
    "capital de solvabilite requis",
    "capital requis de solvabilite",
    "minimum de capital requis",
    "capital requis",
    "scr",
    "mcr",
    "bscr",
    "orsa",
    "ersa",
    "sfcr",
    "rsr",
    "qrt",
    "gouvernance",
    "systeme de gouvernance",
    "fonction actuarielle",
    "fonction gestion des risques",
    "fonction conformite",
    "fonction audit interne",
    "fit and proper",
    "personne prudente",
    "prudent person principle",
    "formule standard",
    "standard formula",
    "modele interne",
    "module de risque",
    "modules de risque",
    "courbe des taux",
    "technical provisions",
    "flux de tresorerie",
    "bilan prudentiel",
    "valorisation prudentielle",
    "volatility adjustment",
    "matching adjustment",
    "eiopa",
    "acpr",
    "bnb",
    "nbb",
    "directive",
    "reglement delegue",
    "delegated regulation",
}

SOLVENCY_DOMAIN_TERMS = {
    "risque de marche",
    "market risk",
    "risque de souscription vie",
    "life underwriting risk",
    "risque de souscription non-vie",
    "non-life underwriting risk",
    "risque de souscription sante",
    "health underwriting risk",
    "risque de contrepartie",
    "risque de defaut de contrepartie",
    "counterparty default risk",
    "risque operationnel",
    "operational risk",
    "risque incorporel",
    "intangible asset risk",
    "risque de taux",
    "risque de taux d'interet",
    "interest rate risk",
    "risque actions",
    "equity risk",
    "risque immobilier",
    "property risk",
    "risque de spread",
    "spread risk",
    "risque de change",
    "currency risk",
    "risque de concentration",
    "market risk concentration",
    "risque de mortalite",
    "mortality risk",
    "risque de longevite",
    "longevity risk",
    "risque d'invalidite",
    "risque de morbidite",
    "disability-morbidity risk",
    "risque de rachat",
    "lapse risk",
    "risque de depenses",
    "expense risk",
    "risque de revision",
    "revision risk",
    "risque catastrophe vie",
    "life catastrophe risk",
    "risque de prime",
    "risque de reserve",
    "risque de primes et reserves",
    "premium risk",
    "reserve risk",
    "risque catastrophe",
    "catastrophe risk",
}

SOLVENCY_DOMAIN_CONTEXT_TERMS = {
    "capital",
    "capital requis",
    "scr",
    "bscr",
    "formule standard",
    "standard formula",
    "module",
    "modules",
    "sous-module",
    "sous-modules",
    "sub-module",
    "sub-modules",
}

OUT_OF_SCOPE_ANSWER = (
    "Hors contexte : la question ne semble pas relever du périmètre Solvabilité II "
    "ou du corpus réglementaire chargé. Reformulez avec un sujet Solvabilité II "
    "(par exemple SCR, MCR, ORSA, provisions techniques, gouvernance, SFCR/RSR, "
    "Directive, Actes délégués ou EIOPA) pour obtenir une réponse sourcée."
)


def expand_query(query: str) -> str:
    expanded = [query]
    lowered = query.lower()
    for key, values in QUERY_SYNONYMS.items():
        if key in lowered:
            expanded.extend(values)
    return " ".join(expanded)


def is_solvency_ii_review_query(query: str) -> bool:
    lowered = query.lower()
    return any(term in lowered for term in SOLVENCY_II_REVIEW_QUERY_TERMS)


def is_solvency_ii_review_chunk(chunk: Chunk) -> bool:
    blob = f"{chunk.source_name} {chunk.source_path} {chunk.version or ''} {chunk.text[:500]}".lower()
    return any(term in blob for term in SOLVENCY_II_REVIEW_SOURCE_TERMS)


def normalize_for_scope(text: str) -> str:
    decomposed = unicodedata.normalize("NFKD", text.lower())
    return "".join(char for char in decomposed if not unicodedata.combining(char))


def is_solvency_scope_question(question: str) -> bool:
    normalized = normalize_for_scope(question)
    if is_solvency_ii_review_query(normalized):
        return True
    if re.search(r"\b(?:art\.?|article)\s*\d+[a-z]?\b", normalized):
        return True
    if any(term in normalized for term in SOLVENCY_SCOPE_TERMS):
        return True
    has_domain_term = any(term in normalized for term in SOLVENCY_DOMAIN_TERMS)
    has_domain_context = any(term in normalized for term in SOLVENCY_DOMAIN_CONTEXT_TERMS)
    return has_domain_term and has_domain_context


def tokenize_for_bm25(text: str) -> list[str]:
    # Unicode-aware tokenization: keeps French accents and regulatory expressions such as 2009/138/CE.
    tokens = re.findall(r"[^\W_]+(?:[-/][^\W_]+)*", text.lower(), flags=re.UNICODE)
    return [token for token in tokens if token not in FRENCH_STOPWORDS and len(token) > 1]


def normalize_retrieval_scores(scores_by_id: dict[str, float]) -> dict[str, float]:
    if not scores_by_id:
        return {}
    values = list(scores_by_id.values())
    low = min(values)
    high = max(values)
    if high == low:
        return {chunk_id: 1.0 for chunk_id in scores_by_id}
    return {
        chunk_id: round((score - low) / (high - low), 3)
        for chunk_id, score in scores_by_id.items()
    }


class SolvencyRetriever:
    def __init__(self, chunks_path: Path = CHUNKS_PATH, use_reranker: bool = True):
        ensure_indexes()
        self.chunks = load_chunks(chunks_path)
        self.by_id = {chunk.chunk_id: chunk for chunk in self.chunks}
        self.bm25 = BM25Okapi([tokenize_for_bm25(chunk.text) for chunk in self.chunks])
        self.last_scores: dict[str, float] = {}
        self.last_raw_scores: dict[str, float] = {}

        self.collection = None
        if chromadb is None or SentenceTransformerEmbeddingFunction is None:
            print("Recherche vectorielle ignorée: chromadb/sentence-transformers indisponible.")
        else:
            try:
                embedding_function = SentenceTransformerEmbeddingFunction(model_name=EMBEDDING_MODEL)
                client = chromadb.PersistentClient(path=str(INDEX_DIR / "chroma"))
                self.collection = client.get_or_create_collection(
                    name="solvency_ii_chunks",
                    embedding_function=embedding_function,
                )
            except Exception as e:
                print(f"Recherche vectorielle ignorée: {type(e).__name__}: {e}")

        self.reranker = None
        if use_reranker and CrossEncoder is not None:
            try:
                self.reranker = CrossEncoder(RERANKER_MODEL)
            except Exception as e:
                print(f"Reranker ignoré: {type(e).__name__}: {e}")

    def bm25_search(self, query: str, k: int = TOP_K_BM25) -> list[str]:
        scores = self.bm25.get_scores(tokenize_for_bm25(expand_query(query)))
        ranked = sorted(enumerate(scores), key=lambda item: item[1], reverse=True)[:k]
        return [self.chunks[index].chunk_id for index, score in ranked if score > 0]

    def vector_search(self, query: str, k: int = TOP_K_VECTOR) -> list[str]:
        if self.collection is None:
            return []
        result = self.collection.query(query_texts=[expand_query(query)], n_results=k)
        return [chunk_id for chunk_id in result.get("ids", [[]])[0] if chunk_id in self.by_id]

    def retrieve(
        self,
        query: str,
        k: int = TOP_K_FINAL,
        document_type: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> list[Chunk]:
        review_query = is_solvency_ii_review_query(query)
        bm25_ids = self.bm25_search(query)
        vector_ids = self.vector_search(query)

        fused = {}
        for rank, chunk_id in enumerate(bm25_ids, start=1):
            fused[chunk_id] = fused.get(chunk_id, 0.0) + BM25_RRF_WEIGHT / (RRF_K + rank)
        for rank, chunk_id in enumerate(vector_ids, start=1):
            fused[chunk_id] = fused.get(chunk_id, 0.0) + VECTOR_RRF_WEIGHT / (RRF_K + rank)

        boosted = {}
        for chunk_id, score in fused.items():
            chunk = self.by_id[chunk_id]
            boost = DOCUMENT_TYPE_BOOSTS.get(chunk.document_type or "Inconnu", 1.0)
            if review_query and is_solvency_ii_review_chunk(chunk):
                boost *= 3.0
            boosted[chunk_id] = score * boost

        candidates = [
            self.by_id[chunk_id]
            for chunk_id, score in sorted(boosted.items(), key=lambda item: item[1], reverse=True)
        ]

        if review_query:
            review_candidates = [
                chunk for chunk in self.chunks
                if is_solvency_ii_review_chunk(chunk) and chunk.chunk_id not in boosted
            ]
            candidates = [
                *[chunk for chunk in candidates if is_solvency_ii_review_chunk(chunk)],
                *[chunk for chunk in candidates if not is_solvency_ii_review_chunk(chunk)],
                *review_candidates[:k],
            ]
        self.last_raw_scores = dict(boosted)
        self.last_scores = normalize_retrieval_scores(self.last_raw_scores)

        if document_type:
            candidates = [c for c in candidates if c.document_type == document_type]
        if jurisdiction:
            candidates = [c for c in candidates if c.jurisdiction == jurisdiction]

        if self.reranker is not None and candidates:
            pairs = [(query, chunk.text) for chunk in candidates]
            scores = self.reranker.predict(pairs)
            ranked_pairs = sorted(zip(scores, candidates), key=lambda item: item[0], reverse=True)
            candidates = [chunk for score, chunk in ranked_pairs]
            if review_query:
                candidates = [
                    *[chunk for chunk in candidates if is_solvency_ii_review_chunk(chunk)],
                    *[chunk for chunk in candidates if not is_solvency_ii_review_chunk(chunk)],
                ]
            self.last_raw_scores = {chunk.chunk_id: float(score) for score, chunk in ranked_pairs}
            self.last_scores = normalize_retrieval_scores(self.last_raw_scores)

        returned = candidates[:k]
        self.last_scores = {chunk.chunk_id: self.last_scores.get(chunk.chunk_id, 0.0) for chunk in returned}
        return returned


# ---- Notebook cell 21 ----
def format_citation(chunk: Chunk, number: int) -> str:
    page = f", page {chunk.page}" if chunk.page else ""
    section = f", {chunk.section}" if chunk.section else ""
    version = f", version/date {chunk.version}" if chunk.version else ""
    temporal = f", temporalité: {chunk.temporal_note}" if getattr(chunk, "temporal_note", None) else ""
    return f"[{number}] {chunk.source_name}{page}{section}{version}{temporal}"


def format_context(chunks: list[Chunk]) -> str:
    blocks = []
    for index, chunk in enumerate(chunks, start=1):
        citation = format_citation(chunk, index)
        temporal = f"\nTEMPORALITÉ: {chunk.temporal_note}" if getattr(chunk, "temporal_note", None) else ""
        blocks.append(f"SOURCE {index}: {citation}{temporal}\n{chunk.text}")
    return "\n\n---\n\n".join(blocks)


SYSTEM_PROMPT = """
Tu es un expert Solvabilité II qui explique la réglementation à des actuaires 
et professionnels de l'assurance qualifiés. Tu combines deux qualités :
la rigueur d'un juriste et la pédagogie d'un consultant senior.

Pour chaque réponse, suis EXACTEMENT cette structure en 4 blocs :

---

## 💡 En clair
Explique le concept en 2-3 phrases simples, comme si tu l'expliquais 
à un collègue intelligent qui découvre le sujet. Pas de jargon inutile. 
Utilise des analogies si elles aident.

## 📐 Ce que dit la réglementation
Donne la définition formelle et technique. Cite les articles exacts 
entre crochets [1], [2]. Mentionne explicitement si c'est la Directive, 
les Actes Délégués, ou les Guidelines EIOPA — la distinction compte. Avec la rigueur pour un actuaire.

## ⚙️ Comment ça marche en pratique
Donne un exemple concret ou une application pratique pour un assureur vie 
ou non-vie belge et français. Si une formule existe dans les sources, cite-la ici.

## ⚠️ Limites et points d'attention
Ce que les sources ne couvrent pas. Ce qu'un superviseur regarderait 
en priorité. Ce qui nécessite un avis expert complémentaire.

---

Règles absolues :
- N'invente aucun chiffre, seuil ou formule absent des sources.
- Si les sources sont insuffisantes sur un point, dis-le explicitement 
  dans "Limites" plutôt que d'improviser.
- Si la question porte sur la revoyure, le review, le réexamen ou l'impact
  2025 de Solvabilité II, réponds d'abord à partir de la Directive (UE) 2025/2
  et des sources EIOPA liées au review. Ne remplace pas ce sujet par les règles
  générales de meilleure estimation, QRT ou fonds propres sauf si la question
  le demande explicitement.
- La section "En clair" ne cite pas d'articles — elle explique.
- La section "Réglementation" cite toujours ses sources.
- Réponds en français.
""".strip()


CITATION_RULES = """
Quand tu cites une source, utilise ce format précis :
- Directive : "selon l'Article 77 §2 de la Directive [1]"
- Actes Délégués : "les Actes Délégués précisent à l'Article 17 [2]"
- EIOPA Guidelines : "les Orientations EIOPA recommandent [3]"
- BNB/ACPR : "la circulaire nationale exige [4]"

Ne dis jamais juste "selon [1]" sans nommer le texte.
La hiérarchie des normes compte :
Directive > Actes Délégués > Guidelines EIOPA > Circulaires nationales.
""".strip()


def _chunk_audit_payload(chunk: Chunk, score: Optional[float] = None) -> dict:
    return {
        "chunk_id": chunk.chunk_id,
        "source_name": chunk.source_name,
        "source_path": chunk.source_path,
        "page": chunk.page,
        "section": chunk.section,
        "document_type": chunk.document_type,
        "jurisdiction": chunk.jurisdiction,
        "version": chunk.version,
        "temporal_note": getattr(chunk, "temporal_note", None),
        "retrieval_score": score,
        "text_preview": normalize_whitespace(chunk.text)[:500],
    }


def audit_log_query(
    question: str,
    chunks: list[Chunk],
    retrieval_scores: Optional[dict[str, float]] = None,
    mode: str = "hybrid",
    use_llm: bool = True,
    model: str = "llama-3.3-70b-versatile",
    answer: str = "",
    scope_status: str = "in_scope",
) -> None:
    retrieval_scores = retrieval_scores or {}
    event = {
        "ts": datetime.now(timezone.utc).isoformat(),
        "user": os.environ.get("USER") or os.environ.get("USERNAME") or "unknown",
        "question": question,
        "mode": mode,
        "scope_status": scope_status,
        "llm_requested": use_llm,
        "model": model if use_llm else None,
        "answer_preview": normalize_whitespace(answer)[:1000],
        "sources": [
            _chunk_audit_payload(chunk, retrieval_scores.get(chunk.chunk_id))
            for chunk in chunks
        ],
    }
    AUDIT_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    with AUDIT_LOG_PATH.open("a", encoding="utf-8") as f:
        f.write(json.dumps(event, ensure_ascii=False) + "\n")


def answer_without_llm(chunks: list[Chunk]) -> str:
    if not chunks or sum(len(chunk.text) for chunk in chunks) < MIN_RETRIEVED_CHARS:
        return "Je ne sais pas à partir des sources fournies."

    excerpts = []
    for index, chunk in enumerate(chunks[:3], start=1):
        snippet = normalize_whitespace(chunk.text).replace("\n", " ")
        if len(snippet) > 320:
            snippet = snippet[:317].rsplit(" ", 1)[0] + "..."
        excerpts.append(f"[{index}] {snippet}")

    return "LLM indisponible pour cette session. Voici les extraits les plus pertinents :\n\n" + "\n\n".join(excerpts)


def ensure_groq_api_key(prompt_if_missing: bool = False) -> bool:
    if os.environ.get("GROQ_API_KEY"):
        return True

    if prompt_if_missing:
        from getpass import getpass
        value = getpass("Enter GROQ_API_KEY: ").strip()
        if value:
            os.environ["GROQ_API_KEY"] = value

    return bool(os.environ.get("GROQ_API_KEY"))


def build_messages(
    question: str,
    chunks: list[Chunk],
    history: Optional[list[dict]] = None,
    extra_instruction: str = "",
) -> list[dict]:
    messages = [
        {
            "role": "system",
            "content": f"{SYSTEM_PROMPT}\n\n{CITATION_RULES}",
        }
    ]

    previous_questions = [
        normalize_whitespace(str(item.get("question", "")))
        for item in (history or [])[-3:]
        if normalize_whitespace(str(item.get("question", "")))
    ]

    history_hint = ""
    if previous_questions:
        history_hint = "Questions précédentes, pour contexte uniquement :\n" + "\n".join(
            f"- {item}" for item in previous_questions
        ) + "\n\n"

    extra_text = ""
    if extra_instruction.strip():
        extra_text = extra_instruction.strip() + "\n\n"

    messages.append(
        {
            "role": "user",
            "content": (
                "LANGUE OBLIGATOIRE : réponds exclusivement en français. "
                "Ne réponds jamais en anglais. "
                "Ta réponse doit commencer exactement par `## 💡 En clair`, "
                "puis suivre les 4 blocs demandés dans le prompt système.\n\n"
                f"{history_hint}"
                f"{extra_text}"
                f"Question : {question}\n\nSources :\n{format_context(chunks)}"
            ),
        }
    )

    return messages


def answer_with_groq(
    question: str,
    chunks: list[Chunk],
    model: str = "llama-3.3-70b-versatile",
    history: Optional[list[dict]] = None,
    extra_instruction: str = "",
) -> str:
    if not chunks or sum(len(chunk.text) for chunk in chunks) < MIN_RETRIEVED_CHARS:
        return "Je ne sais pas à partir des sources fournies."

    if not ensure_groq_api_key(prompt_if_missing=False):
        return answer_without_llm(chunks)

    try:
        from langchain_groq import ChatGroq
    except Exception as e:
        print(f"LLM Groq indisponible: {type(e).__name__}: {e}")
        return answer_without_llm(chunks)

    try:
        llm = ChatGroq(model=model, temperature=0)
        response = llm.invoke(
            build_messages(
                question,
                chunks,
                history=history,
                extra_instruction=extra_instruction,
            )
        )
        return response.content
    except Exception as e:
        error = f"Appel Groq impossible: {type(e).__name__}: {e}"
        print(error)

        fallback = answer_without_llm(chunks)
        if "Voici les extraits les plus pertinents :" in fallback:
            fallback = fallback.split(
                "Voici les extraits les plus pertinents :",
                1,
            )[-1].strip()

        return error + "\n\nVoici les extraits les plus pertinents :\n\n" + fallback


def ask(
    question: str,
    use_llm: bool = True,
    use_reranker: bool = True,
    audience: str = "expert",
    history: Optional[list[dict]] = None,
) -> dict:
    model = "llama-3.3-70b-versatile"
    if not is_solvency_scope_question(question):
        audit_log_query(
            question=question,
            chunks=[],
            retrieval_scores={},
            mode="hybrid",
            use_llm=False,
            model=model,
            answer=OUT_OF_SCOPE_ANSWER,
            scope_status="out_of_scope",
        )
        return {
            "question": question,
            "answer": OUT_OF_SCOPE_ANSWER,
            "chunks": [],
            "retrieval_scores": {},
            "scope_status": "out_of_scope",
        }

    retriever = SolvencyRetriever(use_reranker=use_reranker)
    chunks = retriever.retrieve(question)

    print("Sources retrouvées :")
    for index, chunk in enumerate(chunks, start=1):
        print(format_citation(chunk, index))

    if audience == "vulgarise":
        extra = """
Important : l'utilisateur veut comprendre le concept,
pas mémoriser les articles. Privilégie les analogies
et les exemples concrets. Garde les citations mais
mets-les entre parenthèses en fin de phrase, pas en avant.
""".strip()
    else:
        extra = ""

    if use_llm:
        answer = answer_with_groq(
            question,
            chunks,
            model=model,
            extra_instruction=extra,
            history=history,
        )
    else:
        answer = "LLM désactivé. Inspecte les sources retrouvées ci-dessus."

    audit_log_query(
        question=question,
        chunks=chunks,
        retrieval_scores=retriever.last_scores,
        mode="hybrid",
        use_llm=use_llm,
        model=model,
        answer=answer,
        scope_status="in_scope",
    )

    return {
        "question": question,
        "answer": answer,
        "chunks": chunks,
        "retrieval_scores": retriever.last_scores,
        "scope_status": "in_scope",
    }


# ---- Notebook cell 22 ----
class SolvencyBM25Retriever:
    def __init__(self, chunks_path: Path = CHUNKS_PATH):
        ensure_chunks()
        self.chunks = load_chunks(chunks_path)
        self.bm25 = BM25Okapi([tokenize_for_bm25(chunk.text) for chunk in self.chunks])
        self.last_scores: dict[str, float] = {}
        self.last_raw_scores: dict[str, float] = {}

    def retrieve(self, query: str, k: int = TOP_K_FINAL) -> list[Chunk]:
        scores = self.bm25.get_scores(tokenize_for_bm25(expand_query(query)))
        ranked = sorted(enumerate(scores), key=lambda item: item[1], reverse=True)
        scored_results = [
            (self.chunks[index], float(score))
            for index, score in ranked
            if score > 0
        ][:k]
        self.last_raw_scores = {chunk.chunk_id: score for chunk, score in scored_results}
        self.last_scores = normalize_retrieval_scores(self.last_raw_scores)
        return [chunk for chunk, score in scored_results]


# ---- Notebook cell 23 ----
def ask_bm25(
    question: str,
    use_llm: bool = True,
    audience: str = "expert",
    history: Optional[list[dict]] = None,
) -> dict:
    model = "llama-3.3-70b-versatile"
    if not is_solvency_scope_question(question):
        audit_log_query(
            question=question,
            chunks=[],
            retrieval_scores={},
            mode="bm25",
            use_llm=False,
            model=model,
            answer=OUT_OF_SCOPE_ANSWER,
            scope_status="out_of_scope",
        )
        return {
            "question": question,
            "answer": OUT_OF_SCOPE_ANSWER,
            "chunks": [],
            "retrieval_scores": {},
            "scope_status": "out_of_scope",
        }

    retriever = SolvencyBM25Retriever()
    chunks = retriever.retrieve(question)

    print("Sources retrouvées :")
    for index, chunk in enumerate(chunks, start=1):
        print(format_citation(chunk, index))

    if audience == "vulgarise":
        extra = """
        Important : l'utilisateur veut comprendre le concept,
        pas mémoriser les articles. Privilégie les analogies
        et les exemples concrets. Garde les citations mais
        mets-les entre parenthèses en fin de phrase, pas en avant.
        """
    else:
        extra = ""

    if use_llm:
        answer = answer_with_groq(
            question,
            chunks,
            model=model,
            extra_instruction=extra,
            history=history,
        )
    else:
        answer = "LLM désactivé. Inspecte les sources retrouvées ci-dessus."

    audit_log_query(
        question=question,
        chunks=chunks,
        retrieval_scores=retriever.last_scores,
        mode="bm25",
        use_llm=use_llm,
        model=model,
        answer=answer,
        scope_status="in_scope",
    )

    return {
        "question": question,
        "answer": answer,
        "chunks": chunks,
        "retrieval_scores": retriever.last_scores,
        "scope_status": "in_scope",
    }


# ---- Notebook cell 25 ----
EVAL_QUESTIONS = [
    # Pilier 1 - Provisions techniques
    {
        "question": "Que dit l'Article 77 sur le Best Estimate ?",
        "expected_source_contains": ["meilleure estimation", "flux de trésorerie"],
    },
    {
        "question": "Comment se calcule la Risk Margin ?",
        "expected_source_contains": ["marge de risque", "coût du capital"],
    },
    {
        "question": "Qu'est-ce que la segmentation des engagements selon l'Article 80 ?",
        "expected_source_contains": ["segmentation", "ligne d'activité"],
    },
    {
        "question": "Quelles hypothèses pour les flux de trésorerie futurs ?",
        "expected_source_contains": ["projection", "flux"],
    },

    # Pilier 1 - Capital
    {
        "question": "Définition du SCR selon l'Article 101 ?",
        "expected_source_contains": ["capital de solvabilité requis", "article 101"],
    },
    {
        "question": "Comment se calcule le MCR ?",
        "expected_source_contains": ["minimum de capital", "mcr"],
    },
    {
        "question": "Quels sont les modules de risque de la formule standard ?",
        "expected_source_contains": ["module", "formule standard"],
    },
    {
        "question": "Qu'est-ce que le SCR opérationnel ?",
        "expected_source_contains": ["risque opérationnel"],
    },

    # Pilier 2 - Gouvernance
    {
        "question": "Quelles sont les exigences du système de gouvernance Article 41 ?",
        "expected_source_contains": ["gouvernance", "article 41"],
    },
    {
        "question": "Que prévoit l'Article 45 sur l'ORSA ?",
        "expected_source_contains": ["orsa", "article 45"],
    },
    {
        "question": "Quelles sont les quatre fonctions clés sous SII ?",
        "expected_source_contains": ["fonction", "actuarielle", "audit", "gestion des risques"],
    },
    {
        "question": "Exigences fit and proper Article 42 ?",
        "expected_source_contains": ["compétence", "honorabilité"],
    },

    # Pilier 3 - Reporting
    {
        "question": "Que doit contenir le SFCR ?",
        "expected_source_contains": ["sfcr", "rapport"],
    },
    {
        "question": "Qu'est-ce que le RSR ?",
        "expected_source_contains": ["rsr", "superviseur"],
    },
    {
        "question": "Quels QRTs sont obligatoires ?",
        "expected_source_contains": ["qrt", "états quantitatifs"],
    },

    # Groupes
    {
        "question": "Comment se calcule le SCR groupe ?",
        "expected_source_contains": ["groupe", "consolidé"],
    },
    {
        "question": "Qu'est-ce que la diversification au niveau groupe ?",
        "expected_source_contains": ["diversification", "groupe"],
    },

    # Investissements
    {
        "question": "Principe de la personne prudente Article 132 ?",
        "expected_source_contains": ["personne prudente", "article 132"],
    },

    # Réassurance
    {
        "question": "Comment la réassurance réduit-elle le SCR ?",
        "expected_source_contains": ["réassurance", "atténuation"],
    },

    # Modèle interne
    {
        "question": "Conditions d'approbation d'un modèle interne ?",
        "expected_source_contains": ["modèle interne", "approbation"],
    },
]


def evaluate_retrieval(eval_questions: list[dict] = EVAL_QUESTIONS, k: int = TOP_K_FINAL) -> pd.DataFrame:
    retriever = SolvencyRetriever(use_reranker=False)
    rows = []

    for item in eval_questions:
        question = item["question"]
        expected_terms = [term.lower() for term in item.get("expected_source_contains", [])]
        chunks = retriever.retrieve(question, k=k)

        first_hit_rank = None
        for rank, chunk in enumerate(chunks, start=1):
            blob = f"{chunk.source_name} {chunk.section or ''} {chunk.text}".lower()
            matched_terms = [term for term in expected_terms if term in blob]
            coverage = len(matched_terms) / len(expected_terms) if expected_terms else 1.0
            if coverage >= 0.5:
                first_hit_rank = rank
                break

        rows.append({
            "question": question,
            "expected_terms": expected_terms,
            "match_rule": "term coverage >= 50%",
            f"hit@{k}": first_hit_rank is not None,
            "mrr": 0 if first_hit_rank is None else round(1 / first_hit_rank, 3),
            "first_hit_rank": first_hit_rank,
            "top_sources": "; ".join(format_citation(chunk, i + 1) for i, chunk in enumerate(chunks[:3])),
        })

    return pd.DataFrame(rows)


# Run manually in notebook when needed:
#evaluation = evaluate_retrieval()
#evaluation


# ---- Public runtime loader ----
from functools import lru_cache


def _make_paths_portable() -> None:
    project_dir = Path(__file__).resolve().parent
    index_dir = project_dir / "rag3_index"
    globals().update(
        PROJECT_DIR=project_dir,
        DOCS_DIR=project_dir / "Directive",
        INDEX_DIR=index_dir,
        CHUNKS_PATH=index_dir / "chunks.jsonl",
        MANIFEST_PATH=index_dir / "manifest.json",
        AUDIT_LOG_PATH=project_dir / "audit_log.jsonl",
    )


def _install_fast_hybrid_guards(runtime: dict) -> None:
    """Keep Hybrid queries responsive when the vector index is incomplete."""

    def _chunk_count() -> int:
        chunks_path = runtime.get("CHUNKS_PATH")
        if chunks_path is None or not chunks_path.exists():
            return 0
        return sum(1 for line in chunks_path.read_text(encoding="utf-8").splitlines() if line.strip())

    def _vector_count() -> int:
        chromadb_module = runtime.get("chromadb")
        index_dir = runtime.get("INDEX_DIR")
        if chromadb_module is None or index_dir is None:
            return 0
        chroma_dir = index_dir / "chroma"
        if not chroma_dir.exists():
            return 0
        try:
            client = chromadb_module.PersistentClient(path=str(chroma_dir))
            return client.get_collection("solvency_ii_chunks").count()
        except Exception:
            return 0

    def vector_index_ready() -> bool:
        expected = _chunk_count()
        return expected > 0 and _vector_count() == expected

    original_ask = runtime.get("ask")
    ask_bm25_fn = runtime.get("ask_bm25")
    if not callable(original_ask) or not callable(ask_bm25_fn):
        return

    def ask_fast_hybrid(
        question: str,
        use_llm: bool = True,
        use_reranker: bool = False,
        audience: str = "expert",
        history: list[dict] | None = None,
    ) -> dict:
        if not vector_index_ready():
            result = ask_bm25_fn(
                question,
                use_llm=use_llm,
                audience=audience,
                history=history,
            )
            result["mode_used"] = "BM25 fallback"
            result["hybrid_status"] = f"{_vector_count()}/{_chunk_count()}"
            return result

        return original_ask(
            question,
            use_llm=use_llm,
            use_reranker=use_reranker,
            audience=audience,
            history=history,
        )

    runtime["ask"] = ask_fast_hybrid
    runtime["vector_index_ready"] = vector_index_ready
    globals()["ask"] = ask_fast_hybrid
    globals()["vector_index_ready"] = vector_index_ready


@lru_cache(maxsize=1)
def load_runtime() -> dict:
    _make_paths_portable()
    runtime = globals()
    _install_fast_hybrid_guards(runtime)
    return runtime


def get_callable(name: str):
    return load_runtime()[name]
